#!/usr/bin/env python3
"""Structured Microsoft Word document tool."""

# Public execute() converts validation exceptions into structured tool errors.
# ruff: noqa: TRY003

import os
import tempfile
import zipfile
from typing import Any, ClassVar, cast

from agentuniverse.agent.action.tool.common_tool.file_path_utils import resolve_safe_path
from agentuniverse.agent.action.tool.tool import Tool


class WordDocumentTool(Tool):
    """Create, append, read, and inspect DOCX files using structured blocks."""

    base_dir: str = "."
    max_read_bytes: int = 20 * 1024 * 1024
    max_write_bytes: int = 20 * 1024 * 1024
    max_uncompressed_bytes: int = 100 * 1024 * 1024
    max_archive_entries: int = 5_000
    max_blocks: int = 1_000
    max_text_chars: int = 100_000
    max_table_rows: int = 1_000
    max_table_columns: int = 50

    _BLOCK_TYPES: ClassVar[set[str]] = {"heading", "paragraph", "bullet", "table", "page_break"}
    _METADATA_FIELDS: ClassVar[set[str]] = {"title", "subject", "author", "keywords", "comments", "category"}

    def execute(
        self,
        mode: str,
        file_path: str,
        blocks: list[dict[str, Any]] | None = None,
        overwrite: bool = False,
        template_path: str | None = None,
        metadata: dict[str, str] | None = None,
    ) -> dict[str, Any]:
        try:
            self._validate_config()
            operation = self._mode(mode)
            path = self._path(file_path, "file_path")
            if operation == "create":
                return self._create(path, blocks, overwrite, template_path, metadata)
            if operation == "append":
                return self._append(path, blocks)
            if operation == "read":
                return self._read(path)
            return self._info(path)
        except ImportError as exc:
            return self._error(
                file_path,
                "dependency_error",
                "python-docx is required. Install with: pip install python-docx",
                str(exc),
            )
        except (TypeError, ValueError) as exc:
            return self._error(file_path, "validation_error", str(exc))
        except Exception as exc:
            return self._error(file_path, "operation_error", f"Word operation failed: {exc}")

    @staticmethod
    def _error(path: Any, kind: str, message: str, detail: str | None = None) -> dict[str, Any]:
        result = {"status": "error", "error_type": kind, "error": message, "file_path": path}
        if detail:
            result["detail"] = detail
        return result

    def _validate_config(self) -> None:
        for name in (
            "max_read_bytes",
            "max_write_bytes",
            "max_uncompressed_bytes",
            "max_archive_entries",
            "max_blocks",
            "max_text_chars",
            "max_table_rows",
            "max_table_columns",
        ):
            value = getattr(self, name)
            if isinstance(value, bool) or not isinstance(value, int) or value <= 0:
                raise ValueError(f"{name} must be a positive integer")
        if not isinstance(self.base_dir, str) or not self.base_dir:
            raise ValueError("base_dir must be a non-empty string")

    @staticmethod
    def _mode(mode: str) -> str:
        if not isinstance(mode, str):
            raise TypeError("mode must be a string")
        operation = mode.strip().lower()
        if operation not in {"create", "append", "read", "info"}:
            raise ValueError("mode must be create, append, read, or info")
        return operation

    def _path(self, value: str, field: str) -> str:
        if not isinstance(value, str) or not value:
            raise ValueError(f"{field} must be a non-empty string")
        if os.path.splitext(value)[1].lower() != ".docx":
            raise ValueError(f"{field} must have a .docx extension")
        return cast(str, resolve_safe_path(value, self.base_dir))

    def _check_archive(self, path: str, field: str = "file_path", max_bytes: int | None = None) -> None:
        if not os.path.isfile(path):
            raise ValueError(f"{field} does not exist: {path}")
        size = os.path.getsize(path)
        byte_limit = self.max_read_bytes if max_bytes is None else max_bytes
        if size > byte_limit:
            raise ValueError(f"{field} exceeds its byte limit ({byte_limit})")
        try:
            with zipfile.ZipFile(path) as archive:
                entries = archive.infolist()
                if len(entries) > self.max_archive_entries:
                    raise ValueError(f"{field} exceeds max_archive_entries ({self.max_archive_entries})")
                if sum(item.file_size for item in entries) > self.max_uncompressed_bytes:
                    raise ValueError(f"{field} exceeds max_uncompressed_bytes ({self.max_uncompressed_bytes})")
        except zipfile.BadZipFile as exc:
            raise ValueError(f"{field} is not a valid DOCX archive") from exc

    @staticmethod
    def _document_class() -> Any:
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError("No module named 'docx'") from exc
        return Document

    def _validate_blocks(self, blocks: Any) -> list[dict[str, Any]]:  # noqa: C901
        if not isinstance(blocks, list) or not blocks:
            raise ValueError("blocks must be a non-empty list")
        if len(blocks) > self.max_blocks:
            raise ValueError(f"blocks exceed max_blocks ({self.max_blocks})")
        normalized, total = [], 0
        for index, raw in enumerate(blocks):
            if not isinstance(raw, dict):
                raise TypeError(f"blocks[{index}] must be an object")
            kind = raw.get("type")
            if kind not in self._BLOCK_TYPES:
                raise ValueError(f"blocks[{index}].type is invalid")
            allowed = {"type"}
            block: dict[str, Any] = {"type": kind}
            if kind in {"heading", "paragraph", "bullet"}:
                allowed |= {"text", "style", "level"}
                text = raw.get("text")
                if not isinstance(text, str):
                    raise TypeError(f"blocks[{index}].text must be a string")
                total += len(text)
                block["text"] = text
                style = raw.get("style")
                if style is not None and not isinstance(style, str):
                    raise TypeError(f"blocks[{index}].style must be a string")
                block["style"] = style
                level = raw.get("level", 1 if kind == "heading" else 0)
                upper = 9 if kind == "heading" else 8
                lower = 1 if kind == "heading" else 0
                if isinstance(level, bool) or not isinstance(level, int) or not lower <= level <= upper:
                    raise ValueError(f"blocks[{index}].level must be between {lower} and {upper}")
                block["level"] = level
            elif kind == "table":
                allowed |= {"rows", "style"}
                rows = raw.get("rows")
                if not isinstance(rows, list) or not rows:
                    raise ValueError(f"blocks[{index}].rows must be a non-empty list")
                if len(rows) > self.max_table_rows:
                    raise ValueError("table exceeds max_table_rows")
                width = max((len(row) if isinstance(row, list) else 0) for row in rows)
                if width == 0 or width > self.max_table_columns:
                    raise ValueError("table column count is invalid")
                output_rows = []
                for row in rows:
                    if not isinstance(row, list):
                        raise TypeError("table rows must be lists")
                    output = []
                    for value in row:
                        if value is not None and not isinstance(value, (str, int, float, bool)):
                            raise TypeError("table cells must be scalar values")
                        text = "" if value is None else str(value)
                        total += len(text)
                        output.append(text)
                    output_rows.append(output)
                block.update(rows=output_rows, style=raw.get("style"))
            unknown = set(raw) - allowed
            if unknown:
                raise ValueError(f"blocks[{index}] has unknown fields: {', '.join(sorted(unknown))}")
            normalized.append(block)
        if total > self.max_text_chars:
            raise ValueError(f"block content exceeds max_text_chars ({self.max_text_chars})")
        return normalized

    def _metadata(self, metadata: Any) -> dict[str, str]:
        if metadata is None:
            return {}
        if not isinstance(metadata, dict):
            raise TypeError("metadata must be an object")
        unknown = set(metadata) - self._METADATA_FIELDS
        if unknown:
            raise ValueError(f"metadata has unknown fields: {', '.join(sorted(unknown))}")
        for key, value in metadata.items():
            if not isinstance(value, str) or len(value) > 2_000:
                raise ValueError(f"metadata.{key} must be a string of at most 2000 characters")
        return dict(metadata)

    def _create(self, path: str, blocks: Any, overwrite: bool, template: str | None, metadata: Any) -> dict[str, Any]:
        if not isinstance(overwrite, bool):
            raise TypeError("overwrite must be a boolean")
        if os.path.exists(path) and not overwrite:
            raise ValueError("file exists; set overwrite=true to replace it")
        source = None
        if template is not None:
            source = self._path(template, "template_path")
            self._check_archive(source, "template_path")
        Document = self._document_class()
        document = Document(source) if source else Document()
        validated = self._validate_blocks(blocks)
        self._apply_metadata(document, self._metadata(metadata))
        self._apply_blocks(document, validated)
        self._save(document, path)
        return self._success("create", path, len(validated), document, template_path=source, overwritten=overwrite)

    def _append(self, path: str, blocks: Any) -> dict[str, Any]:
        self._check_archive(path)
        Document = self._document_class()
        document = Document(path)
        validated = self._validate_blocks(blocks)
        self._apply_blocks(document, validated)
        self._save(document, path)
        return self._success("append", path, len(validated), document)

    @staticmethod
    def _apply_metadata(document: Any, metadata: dict[str, str]) -> None:
        for key, value in metadata.items():
            setattr(document.core_properties, key, value)

    @staticmethod
    def _apply_blocks(document: Any, blocks: list[dict[str, Any]]) -> None:
        for block in blocks:
            kind = block["type"]
            if kind == "heading":
                document.add_heading(block["text"], level=block["level"])
            elif kind == "paragraph":
                document.add_paragraph(block["text"], style=block["style"])
            elif kind == "bullet":
                style = block["style"] or (
                    "List Bullet" if block["level"] == 0 else f"List Bullet {min(block['level'] + 1, 3)}"
                )
                document.add_paragraph(block["text"], style=style)
            elif kind == "page_break":
                document.add_page_break()
            else:
                rows, width = block["rows"], max(len(row) for row in block["rows"])
                table = document.add_table(rows=len(rows), cols=width)
                if block.get("style"):
                    table.style = block["style"]
                for row_index, row in enumerate(rows):
                    for column in range(width):
                        table.cell(row_index, column).text = row[column] if column < len(row) else ""

    def _read(self, path: str) -> dict[str, Any]:
        self._check_archive(path)
        document = self._document_class()(path)
        remaining, truncated, blocks = self.max_text_chars, False, []
        for paragraph in document.paragraphs:
            text, remaining, cut = self._bounded(paragraph.text, remaining)
            truncated |= cut
            if text:
                blocks.append(
                    {"type": "paragraph", "text": text, "style": paragraph.style.name if paragraph.style else ""}
                )
        tables = []
        for table in document.tables:
            output = []
            for row in table.rows:
                values = []
                for cell in row.cells:
                    value, remaining, cut = self._bounded(cell.text, remaining)
                    truncated |= cut
                    values.append(value)
                output.append(values)
            tables.append(output)
        return {
            "status": "success",
            "mode": "read",
            "file_path": path,
            "paragraphs": blocks,
            "tables": tables,
            "truncated": truncated,
            "max_text_chars": self.max_text_chars,
        }

    def _info(self, path: str) -> dict[str, Any]:
        self._check_archive(path)
        document = self._document_class()(path)
        props = document.core_properties
        return {
            "status": "success",
            "mode": "info",
            "file_path": path,
            "file_size": os.path.getsize(path),
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "section_count": len(document.sections),
            "metadata": {key: getattr(props, key, "") or "" for key in sorted(self._METADATA_FIELDS)},
        }

    def _save(self, document: Any, path: str) -> None:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        temporary = None
        try:
            with tempfile.NamedTemporaryFile(
                prefix=".word-", suffix=".docx", dir=os.path.dirname(path), delete=False
            ) as output:
                temporary = output.name
            document.save(temporary)
            if os.path.getsize(temporary) > self.max_write_bytes:
                raise ValueError("generated document exceeds max_write_bytes")
            self._check_archive(temporary, "generated document", self.max_write_bytes)
            os.replace(temporary, path)
            temporary = None
        finally:
            if temporary and os.path.exists(temporary):
                os.unlink(temporary)

    @staticmethod
    def _bounded(value: Any, remaining: int) -> tuple[str, int, bool]:
        text = str(value or "").strip()
        if len(text) <= remaining:
            return text, remaining - len(text), False
        if remaining <= 0:
            return "", 0, True
        return text[: max(0, remaining - 1)] + "…", 0, True

    @staticmethod
    def _success(mode: str, path: str, count: int, document: Any, **extra: Any) -> dict[str, Any]:
        return {
            "status": "success",
            "mode": mode,
            "file_path": path,
            "blocks_added": count,
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "file_size": os.path.getsize(path),
            **extra,
        }
